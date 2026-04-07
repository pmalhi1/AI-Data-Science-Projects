"""
tax_gen_master.py
=================
Unified Synthetic Tax Document Generator — Tax Year 2024
---------------------------------------------------------
Generates complete, consistent tax document packages for synthetic persons
using Faker. All forms share the same person profile for data consistency.

Difficulty Levels:
  Level 1 (Easy)          — Salary-only employee, 4–6 supporting docs
  Level 2 (Medium)        — Employee + side business, 7–11 supporting docs
  Level 3 (Moderately Complex) — Multi-income / business owner, 10–15 supporting docs

Mandatory for all levels:
  - Executive Summary (PDF)
  - Prompt Document (DOCX)
  - Form 1040
  - Schedules 1, 2, 3
  - Schedule A, B (always generated; included based on level)
  - Schedule C, D, E, SE (level 2+)
  - Form 4562 (level 2+, if business)
  - Form 8949 + Schedule D (level 2+, if investments)
  - Form 8606 (level 2+, if IRA)
  - 1099-INT, 1099-DIV, W-2 (supporting docs, level-dependent)
  - Other forms via last_form_gen logic (supporting docs)

States shuffled from: California, Texas, New York, Illinois, Florida

REQUIREMENTS:
  pip install reportlab pypdf faker python-docx
"""

from __future__ import annotations

import io
import os
import random
import string
from pathlib import Path
from datetime import date, timedelta

from faker import Faker
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
)
from reportlab.pdfgen import canvas as rl_canvas

# python-docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────────────────
# GLOBAL CONFIG
# ─────────────────────────────────────────────────────────────────────────────

W, H = letter   # 612 × 792
NUM_PERSONS = 2
BASE_OUTPUT = "Generated_Tax_Packages"

TARGET_STATES = {
    "CA": {"name": "California", "cities": ["Los Angeles", "San Francisco", "San Diego", "Sacramento", "Fresno"]},
    "TX": {"name": "Texas",      "cities": ["Houston", "Dallas", "Austin", "San Antonio", "Fort Worth"]},
    "NY": {"name": "New York",   "cities": ["New York City", "Buffalo", "Rochester", "Albany", "Syracuse"]},
    "IL": {"name": "Illinois",   "cities": ["Chicago", "Aurora", "Naperville", "Joliet", "Rockford"]},
    "FL": {"name": "Florida",    "cities": ["Miami", "Orlando", "Tampa", "Jacksonville", "Tallahassee"]},
}

DIFFICULTY_LEVELS = {
    1: {
        "label": "Level 1 (Easy)",
        "description": "Salary-based employee with simple income",
        "income_types": ["w2"],
        "supporting_doc_count": (4, 6),
        "has_business": False,
        "has_investments": False,
        "has_rental": False,
        "has_ira": False,
        "itemize": False,
    },
    2: {
        "label": "Level 2 (Medium)",
        "description": "Employee with side business and investments",
        "income_types": ["w2", "business", "interest", "dividends"],
        "supporting_doc_count": (7, 11),
        "has_business": True,
        "has_investments": True,
        "has_rental": False,
        "has_ira": True,
        "itemize": True,
    },
    3: {
        "label": "Level 3 (Moderately Complex)",
        "description": "Business owner with multiple income streams",
        "income_types": ["w2", "business", "interest", "dividends", "rental", "capital_gains", "foreign"],
        "supporting_doc_count": (10, 15),
        "has_business": True,
        "has_investments": True,
        "has_rental": True,
        "has_ira": True,
        "itemize": True,
    },
}


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — PERSON PROFILE GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def build_person_profile(fake: Faker, person_idx: int, difficulty: int) -> dict:
    """
    Creates a fully consistent synthetic person profile.
    All forms generated for this person draw from this single dict.
    """
    random.seed(person_idx * 1000 + difficulty * 17)
    fake.seed_instance(person_idx * 1000 + difficulty * 17)

    # State selection (shuffle through the 5 target states)
    state_abbr = list(TARGET_STATES.keys())[person_idx % len(TARGET_STATES)]
    state_info = TARGET_STATES[state_abbr]
    city = random.choice(state_info["cities"])

    last_name = fake.last_name()
    first_name = fake.first_name_male() if random.random() < 0.5 else fake.first_name_female()
    spouse_first = fake.first_name_female() if first_name != fake.first_name_female() else fake.first_name_male()

    dob_year = random.randint(1968, 1985)
    dob = date(dob_year, random.randint(1, 12), random.randint(1, 28))

    ssn = fake.ssn()
    spouse_ssn = fake.ssn()
    ein = f"{random.randint(10,99)}-{random.randint(1000000,9999999)}"

    # Income based on difficulty
    d = DIFFICULTY_LEVELS[difficulty]
    wages = round(random.uniform(55000, 130000), 2) if "w2" in d["income_types"] else 0.0
    biz_revenue = round(random.uniform(40000, 180000), 2) if d["has_business"] else 0.0
    biz_expenses = round(biz_revenue * random.uniform(0.35, 0.60), 2) if d["has_business"] else 0.0
    biz_profit = round(biz_revenue - biz_expenses, 2) if d["has_business"] else 0.0

    interest_income = round(random.uniform(200, 3500), 2) if "interest" in d["income_types"] else 0.0
    dividend_income = round(random.uniform(500, 8000), 2) if "dividends" in d["income_types"] else 0.0
    rental_income = round(random.uniform(12000, 48000), 2) if d["has_rental"] else 0.0
    rental_expenses = round(rental_income * random.uniform(0.4, 0.65), 2) if d["has_rental"] else 0.0
    rental_net = round(rental_income - rental_expenses, 2) if d["has_rental"] else 0.0

    capital_gains = round(random.uniform(2000, 35000), 2) if d["has_investments"] else 0.0
    foreign_income = round(random.uniform(3000, 15000), 2) if "foreign" in d["income_types"] else 0.0
    foreign_tax = round(foreign_income * 0.12, 2) if foreign_income else 0.0

    agi_before_adj = wages + biz_profit + interest_income + dividend_income + rental_net + capital_gains + foreign_income
    se_tax_half = round(biz_profit * 0.9235 * 0.153 * 0.5, 2) if d["has_business"] else 0.0
    ira_deduction = round(random.uniform(1000, 7000), 2) if d["has_ira"] else 0.0
    agi = round(agi_before_adj - se_tax_half - ira_deduction, 2)

    # Deductions
    standard_deduction = 29200  # MFJ 2024
    mortgage_interest = round(random.uniform(8000, 22000), 2) if d["itemize"] else 0.0
    state_local_tax = 10000  # SALT cap
    charity = round(random.uniform(500, 5000), 2) if d["itemize"] else 0.0
    itemized_total = mortgage_interest + state_local_tax + charity
    use_itemized = itemized_total > standard_deduction

    deduction_used = round(itemized_total, 2) if use_itemized else standard_deduction
    taxable_income = max(0, round(agi - deduction_used, 2))

    # Simplified tax calculation (MFJ 2024 brackets)
    def calc_tax(inc):
        brackets = [(23200, 0.10), (94300, 0.12), (201050, 0.22),
                    (383900, 0.24), (487450, 0.32), (731200, 0.35), (float('inf'), 0.37)]
        tax, prev = 0.0, 0
        for limit, rate in brackets:
            taxable = min(inc, limit) - prev
            if taxable <= 0: break
            tax += taxable * rate
            prev = limit
        return round(tax, 2)

    income_tax = calc_tax(taxable_income)
    se_tax = round(biz_profit * 0.9235 * 0.153, 2) if d["has_business"] else 0.0
    total_tax = round(income_tax + se_tax, 2)

    fed_withheld = round(wages * random.uniform(0.10, 0.18), 2)
    state_withheld = round(wages * random.uniform(0.03, 0.07), 2)
    est_payments = round(random.uniform(500, 4000), 2) if d["has_business"] else 0.0
    total_payments = round(fed_withheld + est_payments, 2)

    refund_or_owe = round(total_payments - total_tax, 2)

    # Employer / Bank info
    employer_name = fake.company().upper()
    employer_ein = fake.ein()
    employer_address = fake.street_address().upper()

    bank_name = random.choice(["Chase Bank", "Bank of America", "Wells Fargo", "Citibank", "US Bank",
                                "Regions Bank", "TD Bank", "PNC Bank", "SunTrust", "KeyBank"])
    bank_ein = f"{random.randint(10,99)}-{random.randint(1000000,9999999)}"

    # Dependents
    num_dependents = random.randint(0, 3)
    dependents = []
    for _ in range(num_dependents):
        dep_first = fake.first_name()
        dep_age = random.randint(1, 17)
        dep_dob = date(2024 - dep_age, random.randint(1, 12), random.randint(1, 28))
        dependents.append({
            "first": dep_first,
            "last": last_name,
            "full": f"{dep_first} {last_name}",
            "ssn": fake.ssn(),
            "dob": dep_dob.strftime("%m/%d/%Y"),
            "age": dep_age,
            "rel": random.choice(["Son", "Daughter"]),
        })

    child_tax_credit = len([d for d in dependents if d["age"] < 17]) * 2000

    # Business info (level 2+)
    biz_name = f"{last_name} {random.choice(['Consulting', 'Services', 'Solutions', 'Enterprises', 'Group'])} LLC"
    biz_code = random.choice(["541510", "541600", "541511", "531390", "722320", "812990"])
    biz_address = fake.street_address().upper()

    # Rental property info (level 3)
    rental_property_address = f"{fake.street_address()}, {city}, {state_abbr}" if d["has_rental"] else ""

    # IRA info (level 2+)
    ira_balance = round(random.uniform(10000, 120000), 2) if d["has_ira"] else 0.0
    ira_contrib = round(min(7000, ira_deduction), 2) if d["has_ira"] else 0.0

    # Foreign info (level 3)
    foreign_country = random.choice(["United Kingdom", "Germany", "Canada", "France", "Japan", "Australia"]) if foreign_income else ""

    # Brokerage / Investment (level 2+)
    broker_name = random.choice(["Fidelity Investments", "Vanguard", "Charles Schwab", "TD Ameritrade", "E*Trade"])
    broker_ein = f"{random.randint(10,99)}-{random.randint(1000000,9999999)}"
    acct_num = str(random.randint(10000000, 99999999))

    prep_date = date(2025, random.randint(1, 4), random.randint(1, 28)).strftime("%m/%d/%Y")

    return {
        # Identity
        "person_idx": person_idx,
        "difficulty": difficulty,
        "difficulty_label": d["label"],
        "full_name": f"{first_name} {last_name}",
        "first_name": first_name,
        "last_name": last_name,
        "spouse_first": spouse_first,
        "spouse_name": f"{spouse_first} {last_name}",
        "joint_name": f"{first_name} & {spouse_first} {last_name}",
        "ssn": ssn,
        "spouse_ssn": spouse_ssn,
        "ein": ein,
        "dob": dob.strftime("%m/%d/%Y"),
        "occupation": random.choice(["Consultant", "Engineer", "Manager", "Analyst", "Director", "Specialist"]),

        # Address
        "address": fake.street_address(),
        "city": city,
        "state": state_abbr,
        "state_name": state_info["name"],
        "zip": fake.zipcode(),

        # Employer
        "employer_name": employer_name,
        "employer_ein": employer_ein,
        "employer_address": employer_address,
        "employer_city": f"{city}, {state_abbr} {fake.zipcode()}",

        # Bank / Broker
        "bank_name": bank_name,
        "bank_ein": bank_ein,
        "broker_name": broker_name,
        "broker_ein": broker_ein,
        "acct_num": acct_num,

        # Dependents
        "dependents": dependents,
        "num_dependents": num_dependents,

        # Income
        "wages": wages,
        "biz_revenue": biz_revenue,
        "biz_expenses": biz_expenses,
        "biz_profit": biz_profit,
        "interest_income": interest_income,
        "dividend_income": dividend_income,
        "rental_income": rental_income,
        "rental_expenses": rental_expenses,
        "rental_net": rental_net,
        "capital_gains": capital_gains,
        "foreign_income": foreign_income,
        "foreign_tax": foreign_tax,
        "foreign_country": foreign_country,

        # AGI / Tax
        "agi_before_adj": agi_before_adj,
        "se_tax_half": se_tax_half,
        "ira_deduction": ira_deduction,
        "agi": agi,
        "standard_deduction": standard_deduction,
        "mortgage_interest": mortgage_interest,
        "state_local_tax": state_local_tax,
        "charity": charity,
        "itemized_total": itemized_total,
        "use_itemized": use_itemized,
        "deduction_used": deduction_used,
        "taxable_income": taxable_income,
        "income_tax": income_tax,
        "se_tax": se_tax,
        "total_tax": total_tax,
        "child_tax_credit": child_tax_credit,
        "fed_withheld": fed_withheld,
        "state_withheld": state_withheld,
        "est_payments": est_payments,
        "total_payments": total_payments,
        "refund_or_owe": refund_or_owe,

        # Business
        "has_business": d["has_business"],
        "biz_name": biz_name,
        "biz_code": biz_code,
        "biz_address": biz_address,

        # Rental
        "has_rental": d["has_rental"],
        "rental_property_address": rental_property_address,

        # IRA
        "has_ira": d["has_ira"],
        "ira_balance": ira_balance,
        "ira_contrib": ira_contrib,

        # Investments
        "has_investments": d["has_investments"],

        # Meta
        "prep_date": prep_date,
        "filing_status": "Married Filing Jointly",
        "tax_year": "2024",
    }


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — SHARED PDF DRAWING HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def watermark(c):
    c.saveState()
    c.setFont("Helvetica-Bold", 48)
    c.setFillColor(colors.Color(0.85, 0.85, 0.85, alpha=0.35))
    c.translate(W / 2, H / 2)
    c.rotate(40)
    c.drawCentredString(0, 0, "SYNTHETIC TRAINING DATA")
    c.restoreState()

def draw_header_sch(c, title, subtitle, seq, year="2024"):
    c.setFont("Helvetica-Bold", 11)
    c.drawString(36, H - 40, f"SCHEDULE {title}  (Form 1040)  {year}")
    c.setFont("Helvetica", 8)
    c.drawString(36, H - 54, subtitle)
    c.drawRightString(W - 36, H - 40, "OMB No. 1545-0074")
    c.drawRightString(W - 36, H - 52, f"Sequence No. {seq}")
    c.setFont("Helvetica", 7)
    c.drawString(36, H - 66,
        "Department of the Treasury — Internal Revenue Service   "
        "Attach to Form 1040, 1040-SR, or 1040-NR.")
    c.line(36, H - 72, W - 36, H - 72)

def draw_header_form(c, form_num, title, seq, year="2024"):
    c.setFont("Helvetica-Bold", 11)
    c.drawString(36, H - 40, f"Form {form_num}  —  {year}")
    c.setFont("Helvetica", 8)
    c.drawString(36, H - 54, title)
    c.drawRightString(W - 36, H - 40, "OMB No. 1545-0074")
    c.drawRightString(W - 36, H - 52, f"Sequence No. {seq}")
    c.setFont("Helvetica", 7)
    c.drawString(36, H - 66,
        "Department of the Treasury — Internal Revenue Service")
    c.line(36, H - 72, W - 36, H - 72)

def draw_name_ssn(c, name, ssn, y):
    c.setFont("Helvetica", 8)
    c.drawString(36, y, "Name(s) shown on Form 1040")
    c.drawRightString(W - 36, y, "Your social security number")
    c.setFont("Helvetica", 9)
    c.drawString(36, y - 14, name)
    c.drawRightString(W - 36, y - 14, ssn)
    c.line(36, y - 18, W - 36, y - 18)
    return y - 32

def draw_section_title(c, text, y):
    c.setFont("Helvetica-Bold", 9)
    c.drawString(36, y, text)
    c.line(36, y - 4, W - 36, y - 4)
    return y - 18

def draw_line_item(c, line_num, label, value, y, indent=48, bold=False):
    if y < 55:
        return y
    font = "Helvetica-Bold" if bold else "Helvetica"
    c.setFont(font, 8)
    tag = f"{line_num}  {label}" if line_num else f"   {label}"
    max_w = 380
    while c.stringWidth(tag, font, 8) > max_w and len(tag) > 12:
        tag = tag[:-4] + "..."
    c.drawString(indent, y, tag)
    if value:
        c.drawRightString(W - 38, y, str(value))
    c.setStrokeColorRGB(0.6, 0.6, 0.6)
    c.line(W - 140, y - 3, W - 38, y - 3)
    c.setStrokeColorRGB(0, 0, 0)
    return y - 14

def draw_footer(c, form_label):
    c.setFont("Helvetica", 7)
    c.drawString(36, 28, "For Paperwork Reduction Act Notice, see your tax return instructions.")
    c.drawRightString(W - 36, 28, f"{form_label} (2024)")
    c.line(36, 36, W - 36, 36)

def fmt(v):
    if v is None or v == 0:
        return ""
    return f"${v:,.2f}"

def fmtz(v):
    """Format including zero."""
    if v is None:
        return "$0.00"
    return f"${v:,.2f}"

def maybe(v, prob=0.4):
    return v if random.random() < prob else None

def rand_dollar(lo, hi):
    return round(random.uniform(lo, hi), 2)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — FORM 1040
# ─────────────────────────────────────────────────────────────────────────────

def generate_form_1040(p: dict, folder: str):
    path = os.path.join(folder, "Form_1040.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    watermark(c)

    # Header
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(W / 2, H - 45, "Form 1040")
    c.setFont("Helvetica", 10)
    c.drawCentredString(W / 2, H - 60, "U.S. Individual Income Tax Return")
    c.drawCentredString(W / 2, H - 73, "Department of the Treasury — Internal Revenue Service")
    c.setFont("Helvetica-Bold", 10)
    c.drawString(490, H - 45, "2024")
    c.rect(484, H - 58, 70, 20)
    c.line(36, H - 88, W - 36, H - 88)

    # Filing Status
    c.setFont("Helvetica-Bold", 9)
    c.drawString(36, H - 105, "Filing Status")
    c.setFont("Helvetica", 8)
    statuses = ["Single", "Married filing jointly", "Married filing separately",
                "Head of household", "Qualifying surviving spouse"]
    x_pos = 36
    for i, status in enumerate(statuses):
        c.rect(x_pos, H - 125, 8, 8)
        if i == 1:  # MFJ checked
            c.setFont("Helvetica-Bold", 9)
            c.drawString(x_pos + 1, H - 124, "X")
            c.setFont("Helvetica", 8)
        c.drawString(x_pos + 11, H - 122, status)
        x_pos += 115

    c.line(36, H - 138, W - 36, H - 138)

    def draw_field(fx, fy, fw, fh, label, value=""):
        c.setFont("Helvetica", 7)
        c.drawString(fx + 2, fy + fh + 2, label)
        c.rect(fx, fy, fw, fh)
        if value:
            c.setFont("Helvetica", 8)
            c.drawString(fx + 4, fy + 4, str(value))

    y = H - 165
    draw_field(36, y, 256, 16, "Your first name and middle initial", p["first_name"])
    draw_field(300, y, 270, 16, "Last name", p["last_name"])

    y -= 38
    draw_field(36, y, 256, 16, "Spouse's first name and middle initial", p["spouse_first"])
    draw_field(300, y, 270, 16, "Spouse's last name", p["last_name"])

    y -= 38
    draw_field(36, y, 410, 16, "Home address (number and street)", p["address"])
    draw_field(454, y, 118, 16, "Apt. no.", "")

    y -= 38
    draw_field(36, y, 220, 16, "City, town, or post office", p["city"])
    draw_field(264, y, 60, 16, "State", p["state"])
    draw_field(332, y, 100, 16, "ZIP code", p["zip"])
    draw_field(440, y, 132, 16, "Your SSN", p["ssn"])

    # Dependents
    y -= 48
    c.setFont("Helvetica-Bold", 9)
    c.drawString(36, y, "Dependents")
    c.line(36, y - 5, W - 36, y - 5)
    y -= 18
    c.setFont("Helvetica", 7)
    c.drawString(36, y, "(1) Name")
    c.drawString(200, y, "(2) SSN")
    c.drawString(310, y, "(3) Relationship")
    c.drawString(420, y, "(4) Child tax credit")
    y -= 12
    for dep in p["dependents"][:4]:
        c.setFont("Helvetica", 8)
        c.drawString(36, y, dep["full"][:24])
        c.drawString(200, y, dep["ssn"])
        c.drawString(310, y, dep["rel"])
        if dep["age"] < 17:
            c.drawString(420, y, "✓")
        y -= 14

    # Income Section
    y -= 10
    c.setFont("Helvetica-Bold", 10)
    c.drawString(36, y, "Income")
    c.line(36, y - 5, W - 36, y - 5)
    y -= 18

    income_lines = [
        ("1a",  "Wages, salaries, tips from W-2",                   fmtz(p["wages"])),
        ("2b",  "Taxable interest",                                  fmt(p["interest_income"])),
        ("3b",  "Ordinary dividends",                               fmt(p["dividend_income"])),
        ("4b",  "IRA distributions — Taxable amount",               ""),
        ("5b",  "Pensions and annuities — Taxable amount",          ""),
        ("7",   "Capital gain or (loss)",                           fmt(p["capital_gains"])),
        ("8",   "Additional income from Schedule 1, line 10",       fmt(p["biz_profit"] + p["rental_net"])),
        ("11",  "Adjusted Gross Income (AGI)",                      fmtz(p["agi"])),
    ]
    for ln, desc, val in income_lines:
        c.setFont("Helvetica-Bold", 8)
        c.drawString(38, y, ln)
        c.setFont("Helvetica", 8)
        c.drawString(62, y, desc)
        if val:
            c.rect(W - 152, y - 3, 112, 14)
            c.drawRightString(W - 44, y + 2, val)
        c.line(36, y - 8, W - 36, y - 8)
        y -= 20

    # Deductions
    y -= 6
    c.setFont("Helvetica-Bold", 10)
    c.drawString(36, y, "Deductions")
    c.line(36, y - 5, W - 36, y - 5)
    y -= 18

    deduction_label = "Itemized deductions (Schedule A)" if p["use_itemized"] else "Standard deduction"
    ded_lines = [
        ("12", deduction_label,                                    fmtz(p["deduction_used"])),
        ("15", "Taxable income",                                   fmtz(p["taxable_income"])),
    ]
    for ln, desc, val in ded_lines:
        c.setFont("Helvetica-Bold", 8)
        c.drawString(38, y, ln)
        c.setFont("Helvetica", 8)
        c.drawString(62, y, desc)
        c.rect(W - 152, y - 3, 112, 14)
        c.drawRightString(W - 44, y + 2, val)
        c.line(36, y - 8, W - 36, y - 8)
        y -= 20

    # Tax and Credits
    y -= 6
    c.setFont("Helvetica-Bold", 10)
    c.drawString(36, y, "Tax and Credits")
    c.line(36, y - 5, W - 36, y - 5)
    y -= 18

    tax_lines = [
        ("16", "Tax (see instructions)",                           fmtz(p["income_tax"])),
        ("19", "Child tax credit / credit for other dependents",   fmt(p["child_tax_credit"])),
        ("23", "Other taxes from Schedule 2, line 21",             fmt(p["se_tax"])),
        ("24", "Total tax",                                        fmtz(p["total_tax"])),
    ]
    for ln, desc, val in tax_lines:
        c.setFont("Helvetica-Bold", 8)
        c.drawString(38, y, ln)
        c.setFont("Helvetica", 8)
        c.drawString(62, y, desc)
        if val:
            c.rect(W - 152, y - 3, 112, 14)
            c.drawRightString(W - 44, y + 2, val)
        c.line(36, y - 8, W - 36, y - 8)
        y -= 20

    # Payments
    y -= 6
    c.setFont("Helvetica-Bold", 10)
    c.drawString(36, y, "Payments")
    c.line(36, y - 5, W - 36, y - 5)
    y -= 18

    pay_lines = [
        ("25a", "Federal income tax withheld from W-2",            fmtz(p["fed_withheld"])),
        ("26",  "2024 estimated tax payments",                     fmt(p["est_payments"])),
        ("33",  "Total payments",                                  fmtz(p["total_payments"])),
    ]
    for ln, desc, val in pay_lines:
        c.setFont("Helvetica-Bold", 8)
        c.drawString(38, y, ln)
        c.setFont("Helvetica", 8)
        c.drawString(62, y, desc)
        c.rect(W - 152, y - 3, 112, 14)
        c.drawRightString(W - 44, y + 2, val)
        c.line(36, y - 8, W - 36, y - 8)
        y -= 20

    # Refund / Owe
    y -= 6
    c.setFont("Helvetica-Bold", 10)
    c.drawString(36, y, "Refund or Amount Owed")
    c.line(36, y - 5, W - 36, y - 5)
    y -= 18
    if p["refund_or_owe"] >= 0:
        label35 = f"Refund   {fmtz(p['refund_or_owe'])}"
        label37 = "Amount owed   —"
    else:
        label35 = "Refund   —"
        label37 = f"Amount owed   {fmtz(abs(p['refund_or_owe']))}"

    for ln, desc in [("35a", label35), ("37", label37)]:
        c.setFont("Helvetica-Bold", 8)
        c.drawString(38, y, ln)
        c.setFont("Helvetica", 8)
        c.drawString(62, y, desc)
        c.line(36, y - 8, W - 36, y - 8)
        y -= 16

    # Signature
    if y > 80:
        y -= 8
        c.line(36, y, W - 36, y)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(36, y - 12, "Sign Here")
        c.setFont("Helvetica", 7)
        c.drawString(36, y - 22, "Under penalties of perjury, I declare this return is true, correct, and complete.")
        c.rect(36, y - 55, 200, 16)
        c.setFont("Helvetica", 7)
        c.drawString(38, y - 48, "Your signature")
        c.rect(244, y - 55, 100, 16)
        c.drawString(246, y - 48, "Date")
        c.rect(352, y - 55, 220, 16)
        c.drawString(354, y - 48, "Occupation: " + p["occupation"])

    c.setFont("Helvetica", 7)
    c.drawString(36, 28, "Form 1040 (2024)")
    c.drawCentredString(W / 2, 28, "For Disclosure, Privacy Act, and Paperwork Reduction Act Notice, see instructions.")

    c.save()
    print(f"    ✓ Form_1040.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — SCHEDULES 1, 2, 3
# ─────────────────────────────────────────────────────────────────────────────

def generate_schedule_1(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_1.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "1", "Additional Income and Adjustments to Income", "01")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    y = draw_section_title(c, "Part I    Additional Income", y)

    lines = [
        ("3",  "Business income or (loss) — Schedule C",          fmt(p["biz_profit"])),
        ("5",  "Rental real estate, royalties, partnerships",      fmt(p["rental_net"])),
        ("7",  "Unemployment compensation",                        ""),
        ("8z", "Other income",                                     ""),
        ("10", "Total additional income (combine lines 1–9)",      fmt(p["biz_profit"] + p["rental_net"])),
    ]
    for ln, lbl, val in lines:
        y = draw_line_item(c, ln, lbl, val or "—", y, bold=(ln == "10"))

    y -= 8
    c.line(36, y, W - 36, y)
    y -= 10
    y = draw_section_title(c, "Part II    Adjustments to Income", y)

    adj_lines = [
        ("15", "Deductible part of self-employment tax",            fmt(p["se_tax_half"])),
        ("20", "IRA deduction",                                     fmt(p["ira_deduction"])),
        ("26", "Total adjustments (combine lines 11–25)",           fmt(p["se_tax_half"] + p["ira_deduction"])),
    ]
    for ln, lbl, val in adj_lines:
        y = draw_line_item(c, ln, lbl, val or "—", y, bold=(ln == "26"))

    draw_footer(c, "Schedule 1 (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_1.pdf")


def generate_schedule_2(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_2.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "2", "Additional Taxes", "02")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    y = draw_section_title(c, "Part I    Tax", y)
    amt_val = maybe(rand_dollar(200, 2000), 0.15)
    y = draw_line_item(c, "1",  "Alternative minimum tax — Form 6251",          fmt(amt_val) or "—", y)
    y = draw_line_item(c, "2",  "Excess advance premium tax credit repayment",  "—", y)
    p1_total = amt_val or 0
    y = draw_line_item(c, "3",  "Add lines 1 and 2 → Form 1040 line 17",        fmtz(p1_total), y, bold=True)

    y -= 8
    c.line(36, y, W - 36, y)
    y -= 10
    y = draw_section_title(c, "Part II    Other Taxes", y)

    se_net = round(p["biz_profit"] * 0.9235, 2) if p["has_business"] else 0
    add_med = maybe(rand_dollar(200, 1500), 0.15) if p["agi"] > 250000 else None
    other_total = round(p["se_tax"] + (add_med or 0), 2)

    y = draw_line_item(c, "4",  "Self-employment tax — Schedule SE",             fmtz(p["se_tax"]), y)
    y = draw_line_item(c, "11", "Additional Medicare Tax — Form 8959",           fmt(add_med) or "—", y)
    y = draw_line_item(c, "21", "Total other taxes → Form 1040 line 23",        fmtz(other_total), y, bold=True)

    draw_footer(c, "Schedule 2 (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_2.pdf")


def generate_schedule_3(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_3.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "3", "Additional Credits and Payments", "03")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    y = draw_section_title(c, "Part I    Nonrefundable Credits", y)

    foreign_credit = round(p["foreign_tax"] * 0.9, 2) if p["foreign_tax"] else None
    child_care = maybe(rand_dollar(300, 1200), 0.35) if p["num_dependents"] > 0 else None
    edu_credit = maybe(rand_dollar(500, 2500), 0.2)
    nr_total = (foreign_credit or 0) + (child_care or 0) + (edu_credit or 0)

    y = draw_line_item(c, "1", "Foreign tax credit — Form 1116",                    fmt(foreign_credit) or "—", y)
    y = draw_line_item(c, "2", "Child and dependent care expenses — Form 2441",     fmt(child_care) or "—", y)
    y = draw_line_item(c, "3", "Education credits — Form 8863",                     fmt(edu_credit) or "—", y)
    y = draw_line_item(c, "8", "Add lines 1–7 → Form 1040 line 20",                fmtz(nr_total), y, bold=True)

    y -= 8
    c.line(36, y, W - 36, y)
    y -= 10
    y = draw_section_title(c, "Part II    Other Payments and Refundable Credits", y)

    net_premium = maybe(rand_dollar(500, 3000), 0.15)
    ref_total = net_premium or 0
    y = draw_line_item(c, "9",  "Net premium tax credit — Form 8962",              fmt(net_premium) or "—", y)
    y = draw_line_item(c, "15", "Add lines 9–14z → Form 1040 line 31",            fmtz(ref_total), y, bold=True)

    draw_footer(c, "Schedule 3 (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_3.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — SCHEDULES A, B
# ─────────────────────────────────────────────────────────────────────────────

def generate_schedule_a(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_A.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "A", "Itemized Deductions", "07")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    # Medical
    y = draw_section_title(c, "Medical and Dental Expenses", y)
    med_total = rand_dollar(2000, 15000)
    agi_floor  = round(p["agi"] * 0.075, 2)
    med_allowed = max(0, round(med_total - agi_floor, 2))
    y = draw_line_item(c, "1", "Medical and dental expenses",             fmt(med_total), y)
    y = draw_line_item(c, "2", "Enter AGI (Form 1040 line 11)",           fmtz(p["agi"]), y)
    y = draw_line_item(c, "3", "Multiply line 2 by 7.5%",                fmtz(agi_floor), y)
    y = draw_line_item(c, "4", "Subtract line 3 from line 1",            fmtz(med_allowed), y, bold=True)
    y -= 4

    # Taxes Paid
    y = draw_section_title(c, "Taxes You Paid", y)
    y = draw_line_item(c, "5a", "State and local income taxes",           fmtz(p["state_local_tax"]), y)
    real_estate = round(p["mortgage_interest"] * 0.15, 2)
    y = draw_line_item(c, "5b", "State and local real estate taxes",      fmt(real_estate), y)
    y = draw_line_item(c, "5e", "SALT cap ($10,000/$5,000 MFS)",          fmtz(min(10000, p["state_local_tax"] + real_estate)), y, bold=True)
    y -= 4

    # Interest
    y = draw_section_title(c, "Interest You Paid", y)
    y = draw_line_item(c, "8a", "Home mortgage interest (Form 1098)",     fmtz(p["mortgage_interest"]), y)
    y = draw_line_item(c, "10", "Add lines 8a through 9",                 fmtz(p["mortgage_interest"]), y, bold=True)
    y -= 4

    # Charity
    y = draw_section_title(c, "Gifts to Charity", y)
    y = draw_line_item(c, "11", "Gifts by cash or check",                 fmtz(p["charity"]), y)
    y = draw_line_item(c, "14", "Add lines 11 through 13",                fmtz(p["charity"]), y, bold=True)
    y -= 4

    # Total
    y = draw_section_title(c, "Total Itemized Deductions", y)
    y = draw_line_item(c, "17", "Total itemized deductions → Form 1040 Sch. A",
                       fmtz(p["itemized_total"]), y, bold=True)

    draw_footer(c, "Schedule A (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_A.pdf")


def generate_schedule_b(p: dict, folder: str, fake: Faker):
    path = os.path.join(folder, "Schedule_B.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "B", "Interest and Ordinary Dividends", "08")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    # Part I Interest
    y = draw_section_title(c, "Part I    Interest", y)
    c.setFont("Helvetica", 7)
    c.drawString(48, y + 2, "List name of payer.")
    y -= 10

    banks = [p["bank_name"]]
    if p["has_investments"]:
        banks.append(p["broker_name"])
    int_amounts = [round(p["interest_income"] * r, 2) for r in ([0.7, 0.3] if len(banks) > 1 else [1.0])]
    for bank, amt in zip(banks, int_amounts):
        y = draw_line_item(c, "1", bank, fmtz(amt), y)

    y = draw_line_item(c, "4", "Taxable interest → Form 1040 line 2b",    fmtz(p["interest_income"]), y, bold=True)
    y -= 8
    c.line(36, y, W - 36, y)
    y -= 10

    # Part II Dividends
    y = draw_section_title(c, "Part II    Ordinary Dividends", y)
    c.setFont("Helvetica", 7)
    c.drawString(48, y + 2, "List name of payer.")
    y -= 10

    brokerages = [p["broker_name"]] if p["has_investments"] else [p["bank_name"]]
    div_portion = p["dividend_income"]
    y = draw_line_item(c, "5", brokerages[0], fmtz(div_portion), y)
    y = draw_line_item(c, "6", "Total ordinary dividends → Form 1040 line 3b", fmtz(p["dividend_income"]), y, bold=True)
    y -= 8

    # Part III
    y = draw_section_title(c, "Part III    Foreign Accounts and Trusts", y)
    has_foreign = bool(p["foreign_country"])
    c.setFont("Helvetica", 8)
    c.drawString(48, y, "7a  Did you have a financial interest over a foreign account?")
    c.drawString(W - 80, y, "Yes" if has_foreign else "No")
    y -= 14

    draw_footer(c, "Schedule B (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_B.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — SCHEDULE C
# ─────────────────────────────────────────────────────────────────────────────

def generate_schedule_c(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_C.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "C", "Profit or Loss From Business (Sole Proprietorship)", "09")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    # Business info
    c.setFont("Helvetica", 8)
    c.drawString(36, y, f"A  Principal business: {p['biz_name']}")
    y -= 14
    c.drawString(36, y, f"B  Business code: {p['biz_code']}")
    y -= 14
    c.drawString(36, y, f"C  Business name: {p['biz_name']}")
    y -= 14
    c.drawString(36, y, f"D  EIN: {p['ein']}")
    y -= 14
    c.drawString(36, y, f"E  Business address: {p['biz_address']}, {p['city']}, {p['state']}")
    y -= 14
    c.line(36, y, W - 36, y)
    y -= 10

    # Part I Income
    y = draw_section_title(c, "Part I    Income", y)
    y = draw_line_item(c, "1",  "Gross receipts or sales",                         fmtz(p["biz_revenue"]), y)
    y = draw_line_item(c, "3",  "Subtract returns and allowances from line 1",     fmtz(p["biz_revenue"]), y)
    y = draw_line_item(c, "7",  "Gross income. Add lines 5 and 6",                 fmtz(p["biz_revenue"]), y)
    y -= 6
    c.line(36, y, W - 36, y)
    y -= 10

    # Part II Expenses
    y = draw_section_title(c, "Part II    Expenses", y)
    exp_items = {
        "Advertising":         maybe(rand_dollar(500, 5000), 0.5),
        "Car and truck":       maybe(rand_dollar(1000, 8000), 0.4),
        "Contract labor":      maybe(rand_dollar(2000, 20000), 0.4),
        "Insurance":           maybe(rand_dollar(500, 4000), 0.5),
        "Office expense":      maybe(rand_dollar(200, 2000), 0.5),
        "Supplies":            maybe(rand_dollar(300, 3000), 0.5),
        "Travel":              maybe(rand_dollar(500, 6000), 0.35),
        "Meals (50%)":         maybe(rand_dollar(200, 2000), 0.4),
        "Utilities":           maybe(rand_dollar(300, 2500), 0.4),
        "Other expenses":      maybe(rand_dollar(500, 5000), 0.35),
    }
    total_exp_items = sum(v for v in exp_items.values() if v)
    # Keep expenses consistent with profile
    scale = p["biz_expenses"] / max(total_exp_items, 1)
    scaled_exp = {k: round(v * scale, 2) for k, v in exp_items.items() if v}

    for line_n, (label, val) in enumerate(scaled_exp.items(), start=8):
        y = draw_line_item(c, str(line_n), label, fmtz(val), y)
        if y < 120:
            break

    y = draw_line_item(c, "28", "Total expenses before home office",               fmtz(p["biz_expenses"]), y, bold=True)
    y = draw_line_item(c, "31", "Net profit or (loss) → Schedule 1 line 3",       fmtz(p["biz_profit"]), y, bold=True)

    draw_footer(c, "Schedule C (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_C.pdf")
    return p["biz_profit"]


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — SCHEDULES D, E, SE
# ─────────────────────────────────────────────────────────────────────────────

def generate_schedule_d(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_D.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "D", "Capital Gains and Losses", "12")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    c.setFont("Helvetica", 7)
    c.drawString(36, y, "Attach to Form 1040. Use Form 8949 to list transactions.")
    y -= 16
    c.line(36, y, W - 36, y)
    y -= 10

    cg = p["capital_gains"]
    st_gl  = round(cg * 0.35, 2)
    lt_gl  = round(cg * 0.65, 2)
    st_proc = round(st_gl + rand_dollar(5000, 30000), 2)
    st_basis = round(st_proc - st_gl, 2)
    lt_proc = round(lt_gl + rand_dollar(10000, 60000), 2)
    lt_basis = round(lt_proc - lt_gl, 2)

    def draw_d_row(c, ln, desc, proc, basis, gl, y, bold=False):
        if y < 60: return y
        font = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font, 8)
        c.drawString(48, y, ln)
        c.drawString(110, y, desc[:35])
        if proc is not None: c.drawString(290, y, fmtz(proc))
        if basis is not None: c.drawString(380, y, fmtz(basis))
        if gl is not None:
            gl_str = fmtz(gl) if gl >= 0 else f"({fmtz(abs(gl))})"
            c.drawRightString(W - 38, y, gl_str)
        c.setStrokeColorRGB(0.7, 0.7, 0.7)
        c.line(36, y - 3, W - 36, y - 3)
        c.setStrokeColorRGB(0, 0, 0)
        return y - 14

    y = draw_section_title(c, "Part I    Short-Term Capital Gains and Losses", y)
    c.setFont("Helvetica-Bold", 7)
    for x, h in [(48,"Line"),(110,"Description"),(290,"Proceeds"),(380,"Basis"),(W-38,"Gain/(Loss)")]:
        if x == W-38: c.drawRightString(x, y, h)
        else: c.drawString(x, y, h)
    c.line(36, y-4, W-36, y-4); y -= 13

    y = draw_d_row(c, "1b", "From Form 8949 Box A", st_proc, st_basis, st_gl, y)
    y = draw_d_row(c, "7",  "Net short-term gain/(loss)", None, None, st_gl, y, bold=True)
    y -= 8; c.line(36, y, W-36, y); y -= 10

    y = draw_section_title(c, "Part II    Long-Term Capital Gains and Losses", y)
    c.setFont("Helvetica-Bold", 7)
    for x, h in [(48,"Line"),(110,"Description"),(290,"Proceeds"),(380,"Basis"),(W-38,"Gain/(Loss)")]:
        if x == W-38: c.drawRightString(x, y, h)
        else: c.drawString(x, y, h)
    c.line(36, y-4, W-36, y-4); y -= 13

    y = draw_d_row(c, "8b", "From Form 8949 Box D", lt_proc, lt_basis, lt_gl, y)
    y = draw_d_row(c, "15", "Net long-term gain/(loss)", None, None, lt_gl, y, bold=True)
    y -= 8; c.line(36, y, W-36, y); y -= 10

    y = draw_section_title(c, "Part III    Summary", y)
    combined = round(st_gl + lt_gl, 2)
    y = draw_line_item(c, "16", "Net capital gain/(loss) → Form 1040 line 7",
                       fmtz(combined) if combined >= 0 else f"({fmtz(abs(combined))})", y, bold=True)

    draw_footer(c, "Schedule D (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_D.pdf")


def generate_schedule_e(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_E.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "E", "Supplemental Income and Loss (Rental, Royalties, Partnerships)", "13")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    y = draw_section_title(c, "Part I    Income or Loss From Rental Real Estate and Royalties", y)

    c.setFont("Helvetica-Bold", 8)
    c.drawString(36, y, f"Property: {p['rental_property_address']}")
    y -= 14

    exp_items = {
        "Advertising":        maybe(rand_dollar(100, 600), 0.4),
        "Insurance":          maybe(rand_dollar(500, 3000), 0.6),
        "Management fees":    maybe(rand_dollar(500, 4000), 0.4),
        "Mortgage interest":  maybe(rand_dollar(3000, 18000), 0.65),
        "Repairs":            maybe(rand_dollar(300, 4000), 0.5),
        "Taxes":              maybe(rand_dollar(1000, 5000), 0.6),
        "Utilities":          maybe(rand_dollar(500, 3500), 0.4),
        "Depreciation":       maybe(rand_dollar(2000, 9000), 0.75),
    }
    total_exp = sum(v for v in exp_items.values() if v)
    scale = p["rental_expenses"] / max(total_exp, 1)
    scaled_exp = {k: round(v * scale, 2) for k, v in exp_items.items() if v}

    y = draw_line_item(c, "3",  "Rents received",                                  fmtz(p["rental_income"]), y)
    for label, val in scaled_exp.items():
        y = draw_line_item(c, "", label,                                            fmtz(val), y)
        if y < 130: break
    y = draw_line_item(c, "20", "Total expenses",                                  fmtz(p["rental_expenses"]), y, bold=True)
    y = draw_line_item(c, "21", "Net income or (loss) → Schedule 1 line 5",       fmtz(p["rental_net"]), y, bold=True)

    draw_footer(c, "Schedule E (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_E.pdf")


def generate_schedule_se(p: dict, folder: str):
    path = os.path.join(folder, "Schedule_SE.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_sch(c, "SE", "Self-Employment Tax", "17")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    net_earnings = round(p["biz_profit"] * 0.9235, 2)
    ss_subject   = min(net_earnings, 168600)
    ss_tax       = round(ss_subject * 0.124, 2)
    med_tax      = round(net_earnings * 0.029, 2)
    se_tax       = round(ss_tax + med_tax, 2)
    se_deduction = round(se_tax * 0.5, 2)

    y = draw_section_title(c, "Part I    Self-Employment Tax", y)

    lines = [
        ("2",  "Net profit from Schedule C, line 31",             fmtz(p["biz_profit"])),
        ("3",  "Combine lines 1a, 1b, and 2",                     fmtz(p["biz_profit"])),
        ("4a", "Multiply line 3 by 92.35%",                       fmtz(net_earnings)),
        ("4c", "Net earnings from self-employment",                fmtz(net_earnings)),
        ("10", "Multiply SS portion by 12.4%",                    fmtz(ss_tax)),
        ("11", "Multiply line 6 by 2.9%",                         fmtz(med_tax)),
        ("12", "Self-employment tax → Schedule 2 line 4",         fmtz(se_tax)),
        ("13", "Deduction — one-half of SE tax → Sch 1 line 15", fmtz(se_deduction)),
    ]
    for ln, lbl, val in lines:
        y = draw_line_item(c, ln, lbl, val, y, bold=(ln in ("12", "13")))
        if y < 80: break

    y -= 8
    c.setFont("Helvetica-Bold", 8)
    c.drawString(36, y, f"SE Tax: {fmtz(se_tax)}   SE Deduction: {fmtz(se_deduction)}")

    draw_footer(c, "Schedule SE (Form 1040)")
    c.save()
    print(f"    ✓ Schedule_SE.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — FORM 4562, FORM 8949, FORM 8606
# ─────────────────────────────────────────────────────────────────────────────

def generate_form_4562(p: dict, folder: str):
    path = os.path.join(folder, "Form_4562.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_form(c, "4562", "Depreciation and Amortization (Including Information on Listed Property)", "179")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    c.setFont("Helvetica", 8)
    c.drawString(36, y, f"Business: {p['biz_name']}")
    y -= 16
    c.line(36, y, W - 36, y)
    y -= 10

    y = draw_section_title(c, "Part I    Election to Expense Certain Property Under Section 179", y)

    sec179_max   = 1220000
    total_cost   = rand_dollar(5000, min(p["biz_expenses"] * 0.3, 200000))
    elected_cost = round(total_cost * random.uniform(0.5, 1.0), 2)
    sec179_ded   = min(elected_cost, rand_dollar(elected_cost * 0.8, elected_cost))

    lines = [
        ("1",  "Maximum amount",                              f"${sec179_max:,.2f}"),
        ("2",  "Total cost of section 179 property",         fmtz(total_cost)),
        ("7",  "Total elected cost of section 179 property", fmtz(elected_cost)),
        ("12", "Section 179 expense deduction",              fmtz(sec179_ded)),
    ]
    for ln, lbl, val in lines:
        y = draw_line_item(c, ln, lbl, val, y, bold=(ln == "12"))

    y -= 6
    c.line(36, y, W - 36, y)
    y -= 10

    y = draw_section_title(c, "Part II    Special Depreciation and Other Depreciation", y)
    bonus_dep = maybe(rand_dollar(1000, 30000), 0.5)
    y = draw_line_item(c, "14", "Special depreciation allowance (bonus depreciation)", fmt(bonus_dep) or "—", y)
    y -= 6

    y = draw_section_title(c, "Part III    MACRS Depreciation", y)
    macrs_classes = [
        ("5-year property",  "5 yrs",  "HY", "200DB", rand_dollar(5000, 40000)),
        ("7-year property",  "7 yrs",  "HY", "200DB", rand_dollar(3000, 25000)),
        ("39-year nonresid.","39 yrs", "MM", "S/L",   rand_dollar(50000, 200000)),
    ]
    total_macrs = 0
    c.setFont("Helvetica-Bold", 7)
    c.drawString(48, y, "Class"); c.drawString(200, y, "Basis"); c.drawString(280, y, "Life")
    c.drawString(330, y, "Conv."); c.drawString(390, y, "Method"); c.drawRightString(W-38, y, "Deduction")
    c.line(36, y-4, W-36, y-4); y -= 13

    for cls, life, conv, meth, basis in macrs_classes:
        if random.random() < 0.5:
            rate = {"200DB": 0.2, "S/L": 0.026}.get(meth, 0.14)
            deduct = round(basis * rate, 2)
            total_macrs += deduct
            c.setFont("Helvetica", 7)
            c.drawString(48, y, cls); c.drawString(200, y, fmtz(basis))
            c.drawString(280, y, life); c.drawString(330, y, conv)
            c.drawString(390, y, meth); c.drawRightString(W-38, y, fmtz(deduct))
            c.setStrokeColorRGB(0.8,0.8,0.8); c.line(36,y-3,W-36,y-3); c.setStrokeColorRGB(0,0,0)
            y -= 13

    y -= 6
    y = draw_section_title(c, "Part IV    Summary", y)
    total_dep = round(sec179_ded + (bonus_dep or 0) + total_macrs, 2)
    y = draw_line_item(c, "22", "Total depreciation. Add lines 12, 14–17, 19–21", fmtz(total_dep), y, bold=True)

    draw_footer(c, "Form 4562")
    c.save()
    print(f"    ✓ Form_4562.pdf")


def generate_form_8949(p: dict, folder: str, fake: Faker):
    path = os.path.join(folder, "Form_8949.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_form(c, "8949", "Sales and Other Dispositions of Capital Assets", "12B")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    c.setFont("Helvetica", 7)
    c.drawString(36, y, "Before checking Box A, B, or C, see Form(s) 1099-B from your broker.")
    y -= 18

    cg = p["capital_gains"]
    st_gl = round(cg * 0.35, 2)
    lt_gl = round(cg * 0.65, 2)

    companies = [fake.company() for _ in range(6)]

    def draw_8949_part(c, part_label, part_title, checkbox_label, transactions_data, y):
        y = draw_section_title(c, f"{part_label}    {part_title}", y)
        c.setFont("Helvetica", 8)
        c.drawString(48, y, checkbox_label); c.rect(W - 80, y - 1, 10, 10)
        c.setFont("Helvetica-Bold", 8); c.drawString(W - 68, y, "✓")
        y -= 18
        c.setFont("Helvetica-Bold", 7)
        for x, h in [(48,"(a) Desc"),(165,"(b) Acq"),(220,"(c) Sold"),
                     (275,"(d) Proceeds"),(345,"(e) Basis"),(415,"(f)"),(445,"(g)"),(W-38,"(h) Gain/Loss")]:
            if x == W-38: c.drawRightString(x, y, h)
            else: c.drawString(x, y, h)
        c.line(36, y-4, W-36, y-4); y -= 13

        tot_p = tot_b = tot_gl = 0
        for desc, acq, sold, proc, basis, gl in transactions_data:
            if y < 80: break
            c.setFont("Helvetica", 7)
            c.drawString(48, y, desc[:18]); c.drawString(165, y, acq); c.drawString(220, y, sold)
            c.drawString(275, y, fmtz(proc)); c.drawString(345, y, fmtz(basis))
            c.drawString(415, y, "—"); c.drawString(445, y, "—")
            gl_s = fmtz(gl) if gl >= 0 else f"({fmtz(abs(gl))})"
            c.drawRightString(W-38, y, gl_s)
            c.setStrokeColorRGB(0.8,0.8,0.8); c.line(36,y-3,W-36,y-3); c.setStrokeColorRGB(0,0,0)
            y -= 11
            tot_p += proc; tot_b += basis; tot_gl += gl

        c.setFont("Helvetica-Bold", 7)
        c.drawString(48, y, "TOTALS")
        c.drawString(275, y, fmtz(round(tot_p,2))); c.drawString(345, y, fmtz(round(tot_b,2)))
        tot_s = fmtz(round(tot_gl,2)) if tot_gl >= 0 else f"({fmtz(abs(round(tot_gl,2)))})"
        c.drawRightString(W-38, y, tot_s)
        c.line(36, y-4, W-36, y-4); y -= 14
        return y, round(tot_p,2), round(tot_b,2), round(tot_gl,2)

    def make_transaction(term, target_gl_contribution):
        proc = rand_dollar(2000, 40000)
        gl   = round(target_gl_contribution * random.uniform(0.8, 1.2), 2)
        basis = round(proc - gl, 2)
        if term == "short":
            acq_date  = date(2024, random.randint(1, 6), random.randint(1, 28))
            sold_date = date(2024, random.randint(7, 12), random.randint(1, 28))
        else:
            acq_date  = date(random.randint(2019, 2022), random.randint(1, 12), random.randint(1, 28))
            sold_date = date(2024, random.randint(1, 12), random.randint(1, 28))
        company = random.choice(companies)
        return (f"{random.randint(10,200)} sh {company[:15]}",
                acq_date.strftime("%m/%d/%Y"), sold_date.strftime("%m/%d/%Y"),
                proc, basis, gl)

    num_st = random.randint(2, 4)
    num_lt = random.randint(2, 4)
    st_transactions = [make_transaction("short", st_gl / num_st) for _ in range(num_st)]
    lt_transactions = [make_transaction("long",  lt_gl / num_lt) for _ in range(num_lt)]

    y, _, _, _ = draw_8949_part(c, "Part I", "Short-Term (Held ≤ 1 Year)",
        "Box A: Reported on 1099-B showing basis reported to IRS", st_transactions, y)
    c.showPage(); watermark(c)
    c.setFont("Helvetica-Bold", 9); c.drawString(36, H-30, "Form 8949 (2024) — Page 2")
    c.setFont("Helvetica", 8); c.drawString(36, H-44, p["full_name"].upper())
    c.drawRightString(W-36, H-44, p["ssn"]); c.line(36, H-50, W-36, H-50)
    y2 = H - 68
    draw_8949_part(c, "Part II", "Long-Term (Held > 1 Year)",
        "Box D: Reported on 1099-B showing basis reported to IRS", lt_transactions, y2)

    draw_footer(c, "Form 8949")
    c.save()
    print(f"    ✓ Form_8949.pdf")


def generate_form_8606(p: dict, folder: str):
    path = os.path.join(folder, "Form_8606.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    draw_header_form(c, "8606", "Nondeductible IRAs", "48")
    watermark(c)
    y = draw_name_ssn(c, p["full_name"].upper(), p["ssn"], H - 80)

    c.setFont("Helvetica", 7)
    c.drawString(36, y, "File with Form 1040, 1040-SR, or 1040-NR.")
    y -= 14

    ira_limit = 7000
    nonded_contrib  = round(min(p["ira_contrib"], ira_limit), 2)
    prior_basis     = round(p["ira_balance"] * 0.1, 2) if p["ira_balance"] else 0
    total_basis     = round(nonded_contrib + prior_basis, 2)
    total_ira_value = p["ira_balance"]
    basis_ratio     = round(min(1.0, total_basis / max(1, total_ira_value)), 4)

    y = draw_section_title(c, "Part I    Nondeductible Contributions to Traditional IRAs", y)
    lines = [
        ("1",  "Nondeductible contributions to traditional IRAs for 2024",     fmtz(nonded_contrib)),
        ("2",  "Total basis in traditional IRAs for 2023 and prior years",      fmtz(prior_basis)),
        ("3",  "Add lines 1 and 2 — total basis",                               fmtz(total_basis)),
        ("6",  "Value of all traditional IRAs as of 12/31/2024",                fmtz(total_ira_value)),
        ("8",  "Add lines 6 and 7",                                             fmtz(total_ira_value)),
        ("9",  "Divide line 3 by line 8 (decimal, not over 1.000)",             f"{basis_ratio:.4f}"),
        ("14", "Total IRA basis for 2024 and earlier",                          fmtz(total_basis)),
    ]
    for ln, lbl, val in lines:
        y = draw_line_item(c, ln, lbl, val, y, bold=(ln in ("3", "14")))

    y -= 6
    c.line(36, y, W - 36, y)
    y -= 10

    has_conversion = random.random() < 0.3
    y = draw_section_title(c, "Part II    Conversions From Traditional to Roth IRAs", y)
    if has_conversion:
        conv_amount = rand_dollar(5000, 40000)
        taxable_conv = round(conv_amount * (1 - basis_ratio), 2)
        conv_lines = [
            ("16", "Amount converted to Roth IRA in 2024",           fmtz(conv_amount)),
            ("17", "Basis in amount converted",                       fmtz(round(conv_amount - taxable_conv, 2))),
            ("18", "Taxable amount → Form 1040 line 4b",             fmtz(taxable_conv)),
        ]
        for ln, lbl, val in conv_lines:
            y = draw_line_item(c, ln, lbl, val, y, bold=(ln == "18"))
    else:
        c.setFont("Helvetica", 8)
        c.drawString(48, y, "No Roth conversion in 2024.")
        y -= 14

    draw_footer(c, "Form 8606")
    c.save()
    print(f"    ✓ Form_8606.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — W-2
# ─────────────────────────────────────────────────────────────────────────────

def generate_w2(p: dict, folder: str):
    path = os.path.join(folder, "W2.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    watermark(c)

    ss_wages = min(p["wages"], 168600.0)
    ss_tax   = round(ss_wages * 0.062, 2)
    med_tax  = round(p["wages"] * 0.0145, 2)
    state_tax = p["state_withheld"]

    c.setLineWidth(1.0)
    c.rect(30, 480, 540, 275)
    c.line(290, 480, 290, 755)
    c.line(430, 545, 430, 755)
    for y_line in range(575, 755, 30):
        c.line(290, y_line, 570, y_line)
    c.line(30, 725, 290, 725)
    c.line(30, 695, 290, 695)
    c.line(30, 615, 290, 615)
    c.line(30, 585, 290, 585)
    c.line(30, 545, 570, 545)
    c.line(30, 505, 570, 505)

    c.setFont("Helvetica-Bold", 6.5)
    c.drawString(35, 745, "a Employee's social security number")
    c.drawString(35, 715, "b Employer identification number (EIN)")
    c.drawString(35, 685, "c Employer's name, address, and ZIP code")
    c.drawString(35, 605, "d Control number")
    c.drawString(35, 575, "e Employee's name")
    c.drawString(35, 535, "f Employee's address and ZIP code")
    c.drawString(295, 745, "1 Wages, tips, other comp.")
    c.drawString(435, 745, "2 Federal income tax withheld")
    c.drawString(295, 715, "3 Social security wages")
    c.drawString(435, 715, "4 Social security tax withheld")
    c.drawString(295, 685, "5 Medicare wages and tips")
    c.drawString(435, 685, "6 Medicare tax withheld")
    c.drawString(35, 495,  "15 State")
    c.drawString(70, 495,  "Employer's state ID number")
    c.drawString(210, 495, "16 State wages")
    c.drawString(310, 495, "17 State income tax")

    c.setFont("Courier-Bold", 10)
    c.drawString(40, 730, p["ssn"])
    c.drawString(40, 700, p["employer_ein"])
    c.drawString(40, 590, str(random.randint(100000000, 999999999)))
    c.drawString(40, 555, p["full_name"].upper())

    t1 = c.beginText(40, 670)
    t1.setLeading(11)
    for line in [p["employer_name"], p["employer_address"], p["employer_city"]]:
        t1.textLine(line[:35])
    c.drawText(t1)

    t2 = c.beginText(40, 520)
    t2.setLeading(11)
    for line in [p["address"], f"{p['city']}, {p['state']} {p['zip']}"]:
        t2.textLine(line.upper())
    c.drawText(t2)

    c.drawRightString(425, 730, f"{p['wages']:,.2f}")
    c.drawRightString(565, 730, f"{p['fed_withheld']:,.2f}")
    c.drawRightString(425, 700, f"{ss_wages:,.2f}")
    c.drawRightString(565, 700, f"{ss_tax:,.2f}")
    c.drawRightString(425, 670, f"{p['wages']:,.2f}")
    c.drawRightString(565, 670, f"{med_tax:,.2f}")
    c.drawString(35, 485, p["state"])
    c.drawString(75, 485, f"{p['state']}-{random.randint(1000000, 9999999)}")
    c.drawRightString(300, 485, f"{p['wages']:,.2f}")
    c.drawRightString(400, 485, f"{state_tax:,.2f}")

    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(W / 2, 460, "W-2 Wage and Tax Statement 2024")
    c.setFont("Helvetica", 8)
    c.drawCentredString(W / 2, 445, "Department of the Treasury — Internal Revenue Service")

    c.save()
    print(f"    ✓ W2.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 10 — 1099-INT
# ─────────────────────────────────────────────────────────────────────────────

def generate_1099_int(p: dict, folder: str):
    path = os.path.join(folder, "1099_INT.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)

    c.setLineWidth(1.0)
    c.rect(36, 370, 540, 360)
    c.setLineWidth(0.6)
    c.line(306, 400, 306, 730); c.line(441, 440, 441, 650)
    c.line(171, 600, 171, 640); c.line(36, 640, 306, 640)
    c.line(36, 600, 306, 600);  c.line(36, 560, 306, 560)
    c.line(36, 520, 306, 520);  c.line(36, 480, 306, 480)
    c.line(36, 440, 306, 440);  c.line(36, 400, 576, 400)
    c.line(306, 690, 576, 690); c.line(306, 650, 576, 650)
    c.line(306, 615, 576, 615); c.line(306, 580, 576, 580)
    c.line(306, 545, 576, 545); c.line(306, 475, 576, 475)
    c.line(306, 440, 576, 440)

    c.setFont("Helvetica", 6.5)
    c.drawString(40, 720, "PAYER'S name, street address, city or town, state, ZIP")
    c.drawString(40, 630, "PAYER'S TIN");  c.drawString(175, 630, "RECIPIENT'S TIN")
    c.drawString(40, 590, "RECIPIENT'S name")
    c.drawString(40, 550, "Street address (including apt. no.)")
    c.drawString(40, 510, "City or town, state, and ZIP code")
    c.drawString(40, 470, "FATCA"); c.drawString(40, 462, "requirement")
    c.rect(80, 455, 10, 10)
    c.drawString(115, 470, "Account number")

    c.setFont("Helvetica-Bold", 10)
    c.drawString(310, 710, "Interest Income")
    c.drawString(445, 715, "OMB No. 1545-0112")
    c.setFont("Helvetica-Bold", 12)
    c.drawString(445, 695, "Form 1099-INT")
    c.setFont("Helvetica", 7)
    c.drawString(310, 695, "For calendar year 2024")

    c.setFont("Helvetica-Bold", 6.5)
    c.drawString(310, 680, "1 Interest income")
    c.drawString(310, 640, "2 Early withdrawal penalty")
    c.drawString(441, 640, "3 Interest on U.S. Savings Bonds")
    c.drawString(310, 605, "4 Federal income tax withheld")
    c.drawString(310, 430, "14 Tax-exempt and tax credit bond CUSIP no.")

    c.setFont("Helvetica", 8)
    c.drawString(40, 695, p["bank_name"].upper())
    c.drawString(40, 680, "100 MAIN ST")
    c.drawString(40, 665, f"{p['city'].upper()}, {p['state']} {p['zip']}")
    c.drawString(45, 615, p["bank_ein"])
    c.drawString(180, 615, p["ssn"])
    c.drawString(45, 570, p["full_name"].upper())
    c.drawString(45, 530, p["address"].upper())
    c.drawString(45, 490, f"{p['city'].upper()}, {p['state']} {p['zip']}")
    c.drawString(115, 455, p["acct_num"][:12])

    c.setFont("Helvetica", 10)
    c.drawRightString(570, 660, f"$ {p['interest_income']:,.2f}")
    c.drawString(45, 375, p["state"])
    c.drawString(95, 375, f"{p['state']}-{p['acct_num'][:6]}")

    c.setFont("Helvetica-Bold", 10)
    c.drawString(500, 425, "Copy B")
    c.drawString(500, 413, "For Recipient")
    c.setFont("Helvetica", 7)
    c.drawString(40, 430, "This is important tax information and is being furnished to the IRS.")
    c.drawString(36, 355, "Form 1099-INT (Rev. 1-2024)")
    c.drawString(250, 355, "www.irs.gov/Form1099INT")
    c.drawString(390, 355, "Department of the Treasury - Internal Revenue Service")

    c.save()
    print(f"    ✓ 1099_INT.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 11 — 1099-DIV (simplified inline — adapted from blank_1099.py)
# ─────────────────────────────────────────────────────────────────────────────

def generate_1099_div(p: dict, folder: str):
    path = os.path.join(folder, "1099_DIV.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)

    # Outer border
    c.setLineWidth(1.0)
    c.rect(9, 9, letter[0] - 18, letter[1] - 18)

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.drawString(350, H - 30, "Form 1099-DIV")
    c.setFont("Helvetica", 8)
    c.drawString(350, H - 42, "(Rev. January 2024)  •  For calendar year 2024")
    c.drawString(350, H - 54, "Dividends and Distributions")
    c.setFont("Helvetica", 6)
    c.drawString(350, H - 65, "OMB No. 1545-0110")

    mid = 298
    c.setLineWidth(0.5)
    c.line(mid, H - 20, mid, H - 240)
    c.line(9, H - 120, mid, H - 120)
    c.line(9, H - 140, mid, H - 140)
    c.line(9, H - 160, mid, H - 160)
    c.line(9, H - 180, mid, H - 180)
    c.line(9, H - 200, mid, H - 200)
    c.line(9, H - 220, mid, H - 220)
    c.line(mid, H - 90, letter[0] - 9, H - 90)
    c.line(mid, H - 120, letter[0] - 9, H - 120)
    c.line(mid, H - 150, letter[0] - 9, H - 150)
    c.line(mid, H - 180, letter[0] - 9, H - 180)
    c.line(mid, H - 210, letter[0] - 9, H - 210)
    c.line(mid, H - 240, letter[0] - 9, H - 240)

    c.setFont("Helvetica", 6.5)
    c.drawString(14, H - 34, "PAYER'S name, address, city, state, ZIP")
    c.drawString(14, H - 127, "PAYER'S TIN")
    c.drawString(160, H - 127, "RECIPIENT'S TIN")
    c.drawString(14, H - 147, "RECIPIENT'S name")
    c.drawString(14, H - 167, "Street address (including apt. no.)")
    c.drawString(14, H - 187, "City or town, state, and ZIP code")
    c.drawString(14, H - 207, "Account number")

    c.setFont("Helvetica-Bold", 6.5)
    c.drawString(mid + 4, H - 34, "1a Total ordinary dividends")
    c.drawString(mid + 4, H - 64, "1b Qualified dividends")
    c.drawString(mid + 4, H - 94, "2a Total capital gain distributions")
    c.drawString(mid + 4, H - 124, "4 Federal income tax withheld")
    c.drawString(mid + 4, H - 154, "5 Section 199A dividends")
    c.drawString(mid + 4, H - 184, "7 Foreign tax paid")
    c.drawString(mid + 4, H - 214, "8 Foreign country or U.S. territory")

    c.setFont("Helvetica", 9)
    qualified = round(p["dividend_income"] * 0.7, 2)
    cap_gain_dist = round(p["dividend_income"] * 0.15, 2)
    fed_w = round(p["dividend_income"] * 0.05, 2) if random.random() < 0.3 else 0
    sec199a = round(p["dividend_income"] * 0.1, 2)

    c.drawString(14, H - 55,  p["broker_name"].upper())
    c.drawString(14, H - 65,  "500 MARKET ST")
    c.drawString(14, H - 75,  f"{p['city'].upper()}, {p['state']} {p['zip']}")
    c.drawString(14, H - 115, p["broker_ein"])
    c.drawString(160, H - 115, p["ssn"])
    c.drawString(14, H - 135, p["full_name"].upper())
    c.drawString(14, H - 155, p["address"].upper())
    c.drawString(14, H - 175, f"{p['city'].upper()}, {p['state']} {p['zip']}")
    c.drawString(14, H - 195, p["acct_num"])

    c.drawRightString(letter[0] - 14, H - 52, f"$ {p['dividend_income']:,.2f}")
    c.drawRightString(letter[0] - 14, H - 82, f"$ {qualified:,.2f}")
    c.drawRightString(letter[0] - 14, H - 112, f"$ {cap_gain_dist:,.2f}")
    if fed_w: c.drawRightString(letter[0] - 14, H - 142, f"$ {fed_w:,.2f}")
    c.drawRightString(letter[0] - 14, H - 172, f"$ {sec199a:,.2f}")
    if p["foreign_tax"]:
        c.drawRightString(letter[0] - 14, H - 202, f"$ {p['foreign_tax']:,.2f}")
        c.drawString(mid + 4, H - 222, p["foreign_country"])

    c.setFont("Helvetica-Bold", 9)
    c.drawString(letter[0] - 70, H - 100, "Copy B")
    c.drawString(letter[0] - 70, H - 112, "For Recipient")
    c.setFont("Helvetica", 7)
    c.drawString(9, 20, "Form 1099-DIV (Rev. 1-2024)    www.irs.gov/Form1099DIV    Department of the Treasury - Internal Revenue Service")

    c.save()
    print(f"    ✓ 1099_DIV.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 12 — SUPPORTING DOCUMENTS (last_form_gen style)
# ─────────────────────────────────────────────────────────────────────────────

def _draw_support_field_box(c, x, y, w, h, label, value):
    c.setLineWidth(0.5)
    c.rect(x, y, w, h)
    c.setFont("Helvetica-Bold", 7)
    c.drawString(x + 3, y + h - 10, label.upper())
    c.setFont("Helvetica", 9)
    limit = int(w / 6.5)
    c.drawString(x + 5, y + 5, str(value)[:limit])

def _draw_support_header(c, title, form_id):
    c.setStrokeColor(colors.black)
    c.setLineWidth(2)
    c.rect(30, H - 80, W - 60, 50)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, H - 55, title)
    c.setFont("Helvetica", 10)
    c.drawString(40, H - 70, f"Form/Document ID: {form_id} | Tax Year: 2024")
    c.drawRightString(W - 40, H - 55, "OMB No. 1545-XXXX")

def generate_support_tax_form(p: dict, title: str, form_slug: str, folder: str):
    path = os.path.join(folder, f"{form_slug.replace(' ', '_')}.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    _draw_support_header(c, title, form_slug)
    watermark(c)

    _draw_support_field_box(c, 30, H - 152, 276, 60, "Payer/Issuer Name and Address", p["broker_name"])
    _draw_support_field_box(c, 306, H - 152, 276, 60, "Recipient Name and Address", p["full_name"])
    _draw_support_field_box(c, 30, H - 202, 138, 50, "Payer TIN", p["broker_ein"])
    _draw_support_field_box(c, 168, H - 202, 138, 50, "Recipient TIN", p["ssn"])

    y_start = H - 282
    gross = rand_dollar(1000, 50000)
    _draw_support_field_box(c, 30, y_start, 184, 40, "Box 1: Gross Proceeds", f"${gross:,.2f}")
    _draw_support_field_box(c, 214, y_start, 184, 40, "Box 2: Federal Tax Withheld", "$0.00")
    _draw_support_field_box(c, 398, y_start, 184, 40, "Box 3: Account Number", p["acct_num"])

    c.setFont("Helvetica", 7)
    c.drawString(30, 30, f"{form_slug} (2024)   Synthetic Training Data")
    c.save()
    print(f"    ✓ {form_slug.replace(' ','_')}.pdf")

def generate_statement(p: dict, title: str, folder: str, fake: Faker):
    slug = title.replace(" ", "_").replace("/", "-")
    path = os.path.join(folder, f"{slug}.pdf")
    c = rl_canvas.Canvas(path, pagesize=letter)
    watermark(c)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(30, H - 50, title)
    c.setFont("Helvetica", 10)
    c.drawString(30, H - 65, f"Account Holder: {p['full_name']} | Date: {p['prep_date']}")

    cols = [80, 250, 100, 100]
    y = H - 120
    headers = ["Date", "Description", "Debit", "Credit"]
    curr_x = 30
    for i, hdr in enumerate(headers):
        c.rect(curr_x, y, cols[i], 20)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(curr_x + 5, y + 6, hdr)
        curr_x += cols[i]

    for _ in range(10):
        y -= 20
        if y < 60: break
        row_data = [p["prep_date"], fake.catch_phrase()[:30], f"{random.randint(10, 500)}", "0.00"]
        curr_x = 30
        for i, val in enumerate(row_data):
            c.rect(curr_x, y, cols[i], 20)
            c.setFont("Helvetica", 9)
            c.drawString(curr_x + 5, y + 6, str(val)[:int(cols[i] / 6)])
            curr_x += cols[i]

    c.setFont("Helvetica", 7)
    c.drawString(30, 30, f"{title} (2024)   Synthetic Training Data")
    c.save()
    print(f"    ✓ {slug}.pdf")


def generate_supporting_documents(p: dict, folder: str, fake: Faker, difficulty: int):
    """
    Generates additional supporting documents based on difficulty and income profile.
    Level 1: 4-6 docs  |  Level 2: 7-11 docs  |  Level 3: 10-15 docs
    """
    d = DIFFICULTY_LEVELS[difficulty]
    min_docs, max_docs = d["supporting_doc_count"]

    all_possible_forms = [
        ("1099-B",            "Proceeds From Broker/Barter Exchange",       "tax"),
        ("1099-R",            "Distributions From Pensions/IRA",             "tax"),
        ("1099-NEC",          "Nonemployee Compensation",                    "tax"),
        ("1099-MISC",         "Miscellaneous Income",                        "tax"),
        ("Schedule-K1",       "Partner's Share of Income/Credits",           "tax"),
        ("Form-1095-A",       "Health Insurance Marketplace Statement",      "tax"),
        ("Form-1098",         "Mortgage Interest Statement",                 "tax"),
        ("Property-Tax-Bill", "Property Tax Bill / Assessment",              "statement"),
        ("Brokerage-Summary", "Monthly Investment Summary",                  "statement"),
        ("Rental-Agreement",  "Residential Lease & Expense Summary",         "statement"),
        ("Business-Receipt",  "General Business Expense Receipt",            "statement"),
        ("Bank-Statement",    "Monthly Transaction Report",                  "statement"),
        ("HSA-IRA-Statement", "HSA/IRA Contribution Statement",              "statement"),
        ("K1-Partnership",    "Partnership K-1 Income Statement",            "tax"),
        ("Form-2441",         "Child and Dependent Care Expenses",           "tax"),
    ]

    # Filter based on difficulty/profile
    eligible = []
    for slug, title, doc_type in all_possible_forms:
        if "1099-R" in slug and not p["has_ira"]: continue
        if "K1" in slug and difficulty < 3: continue
        if "Rental" in slug and not p["has_rental"]: continue
        if "Form-1098" in slug and p["mortgage_interest"] == 0: continue
        if "Form-2441" in slug and p["num_dependents"] == 0: continue
        eligible.append((slug, title, doc_type))

    num_to_generate = random.randint(min_docs, max_docs)
    selected = random.sample(eligible, min(num_to_generate, len(eligible)))

    for slug, title, doc_type in selected:
        if doc_type == "statement":
            generate_statement(p, title, folder, fake)
        else:
            generate_support_tax_form(p, title, slug, folder)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 13 — EXECUTIVE SUMMARY (PDF)
# ─────────────────────────────────────────────────────────────────────────────

def generate_executive_summary(p: dict, folder: str):
    path = os.path.join(folder, "Executive_Summary.pdf")

    def get_tbl_style():
        return TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ])

    doc = SimpleDocTemplate(path, pagesize=letter,
                             leftMargin=0.5*inch, rightMargin=0.5*inch)
    styles = getSampleStyleSheet()
    head_style = ParagraphStyle('SH', parent=styles['Heading2'],
                                 fontSize=12, spaceBefore=12, spaceAfter=6)
    elements = []

    # Page 1: Cover
    elements.append(Paragraph(f"Taxpayer Executive Summary — Tax Year {p['tax_year']}", styles['Title']))
    elements.append(Paragraph(f"Client: {p['joint_name']}", styles['Normal']))
    elements.append(Paragraph(f"Date Prepared: {p['prep_date']}", styles['Normal']))
    elements.append(Paragraph(f"Difficulty Level: {p['difficulty_label']}", styles['Normal']))
    elements.append(Spacer(1, 12))

    t1 = Table([
        ["Field", "Detail"],
        ["Primary Taxpayer",     p["full_name"]],
        ["Spouse",               p["spouse_name"]],
        ["Filing Status",        p["filing_status"]],
        ["State of Residence",   p["state_name"]],
        ["Address",              f"{p['address']}, {p['city']}, {p['state']} {p['zip']}"],
        ["Dependents",           f"{p['num_dependents']} dependent(s)" if p["num_dependents"] else "None"],
        ["Primary Occupation",   p["occupation"]],
    ], colWidths=[2.2*inch, 3.8*inch])
    t1.setStyle(get_tbl_style())
    elements.append(Paragraph("1. Client Profile and Filing Status", head_style))
    elements.append(t1)
    elements.append(PageBreak())

    # Page 2: Income Detail
    elements.append(Paragraph("2. Income & AGI Detail", head_style))
    inc_rows = [["Income Source", "Amount", "Notes"]]
    if p["wages"]:
        inc_rows.append(["W-2 Wages", fmtz(p["wages"]), f"From {p['employer_name'][:30]}"])
    if p["biz_profit"]:
        inc_rows.append(["Schedule C Net Profit", fmtz(p["biz_profit"]), f"{p['biz_name'][:30]}"])
    if p["interest_income"]:
        inc_rows.append(["Interest Income (1099-INT)", fmtz(p["interest_income"]), f"{p['bank_name'][:30]}"])
    if p["dividend_income"]:
        inc_rows.append(["Dividends (1099-DIV)", fmtz(p["dividend_income"]), f"{p['broker_name'][:30]}"])
    if p["rental_net"]:
        inc_rows.append(["Rental Net Income (Sch E)", fmtz(p["rental_net"]), "Rental property"])
    if p["capital_gains"]:
        inc_rows.append(["Capital Gains (Sch D)", fmtz(p["capital_gains"]), "From investments"])
    if p["foreign_income"]:
        inc_rows.append(["Foreign Income", fmtz(p["foreign_income"]), p["foreign_country"]])
    inc_rows.append(["SE Tax Deduction (Sch 1)", f"({fmtz(p['se_tax_half'])})" if p["se_tax_half"] else "—", "Half of SE tax"])
    inc_rows.append(["IRA Deduction", f"({fmtz(p['ira_deduction'])})" if p["ira_deduction"] else "—", "Traditional IRA"])
    inc_rows.append(["ADJUSTED GROSS INCOME", fmtz(p["agi"]), "Form 1040, Line 11"])

    t2 = Table(inc_rows, colWidths=[2.5*inch, 1.5*inch, 2.0*inch])
    t2.setStyle(get_tbl_style())
    t2.setStyle(TableStyle([('FONTNAME', (0, len(inc_rows)-1), (-1, len(inc_rows)-1), 'Helvetica-Bold')]))
    elements.append(t2)
    elements.append(PageBreak())

    # Page 3: Tax Outcome
    elements.append(Paragraph("3. Federal & State Tax Outcome", head_style))
    refund_label = "Refund" if p["refund_or_owe"] >= 0 else "Amount Owed"
    refund_val = fmtz(abs(p["refund_or_owe"]))
    t3 = Table([
        ["Metric", "Federal", "State"],
        ["Total Income Tax",           fmtz(p["income_tax"]),     fmtz(p["state_withheld"])],
        ["SE Tax",                     fmtz(p["se_tax"]),         "N/A"],
        ["Child Tax Credit",           fmtz(p["child_tax_credit"]), "—"],
        ["Total Tax Liability",        fmtz(p["total_tax"]),      "—"],
        ["Total Payments / Withheld",  fmtz(p["total_payments"]), fmtz(p["state_withheld"])],
        [refund_label,                 refund_val,                "See state return"],
    ], colWidths=[2.5*inch, 1.75*inch, 1.75*inch])
    t3.setStyle(get_tbl_style())
    elements.append(t3)
    elements.append(PageBreak())

    # Page 4: Remarks
    elements.append(Paragraph("4. Final Remarks", head_style))
    remarks = [
        f"This executive summary provides a clear overview of tax results for {p['joint_name']} for Tax Year 2024.",
        f"Difficulty classification: {p['difficulty_label']} — {DIFFICULTY_LEVELS[p['difficulty']]['description']}.",
        "All income and deduction items reflect the synthetic data generated for training purposes.",
        f"State of residence: {p['state_name']}. State-specific credits and rules apply.",
        "Please consult a licensed tax professional for actual tax preparation.",
    ]
    for r in remarks:
        elements.append(Paragraph(r, styles['Normal']))
        elements.append(Spacer(1, 6))

    doc.build(elements)
    print(f"    ✓ Executive_Summary.pdf")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 14 — PROMPT DOCUMENT (DOCX)
# ─────────────────────────────────────────────────────────────────────────────

def generate_prompt_docx(p: dict, folder: str):
    path = os.path.join(folder, "Tax_Data_Prompt.docx")
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    def add_bullet(doc, text, value=""):
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(text)
        run.font.bold = True
        if value:
            para.add_run(f"  {value}")
        else:
            para.add_run("  ________________________________________________")

    title = doc.add_heading('TAX RETURN DATA COLLECTION PROMPT — 2024', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('1. Personal & Family Details', level=1)
    add_bullet(doc, "Primary Taxpayer Name:", p["full_name"])
    add_bullet(doc, "Spouse Name:", p["spouse_name"])
    add_bullet(doc, "Filing Status:", p["filing_status"])
    add_bullet(doc, "Current Address:", f"{p['address']}, {p['city']}, {p['state']} {p['zip']}")
    add_bullet(doc, "State of Residence:", p["state_name"])
    add_bullet(doc, "Primary SSN:", p["ssn"])
    add_bullet(doc, "Difficulty Level:", p["difficulty_label"])
    for i, dep in enumerate(p["dependents"], 1):
        add_bullet(doc, f"Dependent {i} (Name/SSN/Age/Rel):",
                   f"{dep['full']} / {dep['ssn']} / {dep['age']} / {dep['rel']}")

    doc.add_heading('2. Income Details', level=1)
    add_bullet(doc, "W-2 Employment Income:", fmtz(p["wages"]) if p["wages"] else "N/A")
    add_bullet(doc, "Employer:", p["employer_name"] if p["wages"] else "N/A")
    add_bullet(doc, "Business Income (Schedule C):", fmtz(p["biz_profit"]) if p["has_business"] else "N/A")
    add_bullet(doc, "Business Name:", p["biz_name"] if p["has_business"] else "N/A")
    add_bullet(doc, "Interest Income (1099-INT):", fmtz(p["interest_income"]) if p["interest_income"] else "N/A")
    add_bullet(doc, "Dividend Income (1099-DIV):", fmtz(p["dividend_income"]) if p["dividend_income"] else "N/A")
    add_bullet(doc, "Capital Gains (Schedule D):", fmtz(p["capital_gains"]) if p["has_investments"] else "N/A")
    add_bullet(doc, "Rental Net Income (Schedule E):", fmtz(p["rental_net"]) if p["has_rental"] else "N/A")
    add_bullet(doc, "Foreign Income:", fmtz(p["foreign_income"]) if p["foreign_income"] else "N/A")
    add_bullet(doc, "Foreign Country:", p["foreign_country"] if p["foreign_country"] else "N/A")

    doc.add_heading('3. Deductions & Credits', level=1)
    add_bullet(doc, "Deduction Type:", "Itemized" if p["use_itemized"] else "Standard Deduction")
    add_bullet(doc, "Total Deduction Amount:", fmtz(p["deduction_used"]))
    add_bullet(doc, "Mortgage Interest:", fmtz(p["mortgage_interest"]) if p["mortgage_interest"] else "N/A")
    add_bullet(doc, "State & Local Taxes (SALT):", fmtz(p["state_local_tax"]) if p["use_itemized"] else "N/A")
    add_bullet(doc, "Charitable Contributions:", fmtz(p["charity"]) if p["charity"] else "N/A")
    add_bullet(doc, "Child Tax Credit:", fmtz(p["child_tax_credit"]) if p["child_tax_credit"] else "N/A")
    add_bullet(doc, "IRA Deduction:", fmtz(p["ira_deduction"]) if p["has_ira"] else "N/A")

    doc.add_heading('4. Tax & Payment Summary', level=1)
    add_bullet(doc, "Adjusted Gross Income (AGI):", fmtz(p["agi"]))
    add_bullet(doc, "Taxable Income:", fmtz(p["taxable_income"]))
    add_bullet(doc, "Federal Income Tax:", fmtz(p["income_tax"]))
    add_bullet(doc, "Self-Employment Tax:", fmtz(p["se_tax"]) if p["has_business"] else "N/A")
    add_bullet(doc, "Total Tax Liability:", fmtz(p["total_tax"]))
    add_bullet(doc, "Federal Tax Withheld (W-2):", fmtz(p["fed_withheld"]))
    add_bullet(doc, "Estimated Tax Payments:", fmtz(p["est_payments"]) if p["est_payments"] else "N/A")
    add_bullet(doc, "Total Payments:", fmtz(p["total_payments"]))
    refund_label = "Refund Amount" if p["refund_or_owe"] >= 0 else "Amount Owed"
    add_bullet(doc, f"{refund_label}:", fmtz(abs(p["refund_or_owe"])))

    doc.add_heading('5. Bank & Investment Account Information', level=1)
    add_bullet(doc, "Primary Bank:", p["bank_name"])
    add_bullet(doc, "Broker / Investment Firm:", p["broker_name"] if p["has_investments"] else "N/A")
    add_bullet(doc, "Account Number (partial):", p["acct_num"][:4] + "****")

    doc.add_heading('6. State-Specific Notes', level=1)
    state_notes = {
        "CA": ["CalEITC eligibility (if income qualifies)", "California SDI withheld",
               "CA Renter's Credit (if applicable)", "Mental Health Services Tax (income > $1M)"],
        "TX": ["No state income tax", "Property tax deduction", "Franchise tax (if business)"],
        "NY": ["New York City resident tax (if NYC)", "NY child tax credit", "STAR property tax credit"],
        "IL": ["Flat 4.95% IL income tax rate", "IL Property Tax Credit", "IL Education Expense Credit"],
        "FL": ["No state income tax", "Property tax homestead exemption", "FL Sales tax considerations"],
    }
    for note in state_notes.get(p["state"], ["See state instructions"]):
        add_bullet(doc, note)

    doc.add_heading('7. Compliance & Prior Year Notes', level=1)
    add_bullet(doc, "Prior-Year Compliance Status:", "In good standing (synthetic)")
    add_bullet(doc, "Estimated Payments Made (Q1–Q4):", fmtz(p["est_payments"]) if p["est_payments"] else "None")
    add_bullet(doc, "IRA Contribution for 2024:", fmtz(p["ira_contrib"]) if p["has_ira"] else "N/A")
    add_bullet(doc, "FBAR / Foreign Account:", "Yes" if p["foreign_country"] else "No")

    doc.save(path)
    print(f"    ✓ Tax_Data_Prompt.docx")


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 15 — MASTER ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_package_for_person(fake: Faker, person_idx: int, difficulty: int):
    d_info = DIFFICULTY_LEVELS[difficulty]
    print(f"\n{'='*65}")
    print(f"  Person {person_idx} | {d_info['label']}")
    print(f"  {d_info['description']}")
    print(f"{'='*65}")

    # Build consistent profile
    p = build_person_profile(fake, person_idx, difficulty)

    folder_name = (
        f"Person_{person_idx}_{p['last_name']}_{p['state']}"
        f"_D{difficulty}"
    )
    folder = os.path.join(BASE_OUTPUT, folder_name)
    os.makedirs(folder, exist_ok=True)

    print(f"\n  Profile: {p['full_name']} ({p['state_name']})")
    print(f"  AGI: {fmtz(p['agi'])}  |  Total Tax: {fmtz(p['total_tax'])}")
    print(f"  Refund/Owe: {fmtz(abs(p['refund_or_owe']))} ({'Refund' if p['refund_or_owe'] >= 0 else 'Owed'})")
    print(f"\n  Generating mandatory documents...")

    # ── MANDATORY DOCUMENTS (all levels) ──────────────────────────────────────
    generate_executive_summary(p, folder)
    generate_prompt_docx(p, folder)
    generate_form_1040(p, folder)
    generate_schedule_1(p, folder)
    generate_schedule_2(p, folder)
    generate_schedule_3(p, folder)
    generate_schedule_a(p, folder)
    generate_schedule_b(p, folder, fake)
    generate_w2(p, folder)
    generate_1099_int(p, folder)

    print(f"\n  Generating level-specific documents ({d_info['label']})...")

    # ── LEVEL 2+ DOCUMENTS ────────────────────────────────────────────────────
    if difficulty >= 2:
        if p["has_business"]:
            generate_schedule_c(p, folder)
            generate_schedule_se(p, folder)
            generate_form_4562(p, folder)
        if p["has_investments"]:
            generate_schedule_d(p, folder)
            generate_form_8949(p, folder, fake)
            generate_1099_div(p, folder)
        if p["has_ira"]:
            generate_form_8606(p, folder)

    # ── LEVEL 3 DOCUMENTS ─────────────────────────────────────────────────────
    if difficulty >= 3:
        if p["has_rental"]:
            generate_schedule_e(p, folder)

    # ── SUPPORTING DOCUMENTS (variable count by difficulty) ───────────────────
    print(f"\n  Generating supporting documents...")
    generate_supporting_documents(p, folder, fake, difficulty)

    total_files = len([f for f in os.listdir(folder)])
    print(f"\n  ✅ {total_files} files generated → {folder}")
    return p


def main():
    os.makedirs(BASE_OUTPUT, exist_ok=True)
    fake = Faker('en_US')

    print("\n" + "="*65)
    print("   SYNTHETIC TAX DOCUMENT GENERATOR — TAX YEAR 2024")
    print("   Generating complete packages for {} person(s)".format(NUM_PERSONS))
    print("="*65)

    # Assign difficulties across persons — cycles through levels
    # You can customize this mapping as needed
    difficulty_assignments = {
        1: 1,   # Person 1 → Level 1 (Easy)
        2: 3,   # Person 2 → Level 3 (Moderately Complex)
        # Add more persons and their levels here
    }

    profiles = []
    for i in range(1, NUM_PERSONS + 1):
        difficulty = difficulty_assignments.get(i, (i % 3) + 1)
        p = generate_package_for_person(fake, i, difficulty)
        profiles.append(p)

    # Print final summary
    print("\n\n" + "="*65)
    print("   GENERATION COMPLETE")
    print("="*65)
    for p in profiles:
        folder_name = f"Person_{p['person_idx']}_{p['last_name']}_{p['state']}_D{p['difficulty']}"
        folder = os.path.join(BASE_OUTPUT, folder_name)
        file_count = len(os.listdir(folder))
        print(f"  • {p['full_name']:30} | {p['difficulty_label']:28} | {file_count:2d} files | {folder}")

    print(f"\n  All packages saved to: ./{BASE_OUTPUT}/\n")


if __name__ == "__main__":
    main()